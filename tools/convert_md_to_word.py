import os
import re
import sys
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup


def get_folder_depth(file_path, base_path):
    relative_path = os.path.relpath(file_path, base_path)
    depth = relative_path.count(os.sep)
    return depth


def get_folder_structure(file_path, base_path):
    relative_path = os.path.relpath(file_path, base_path)
    folders = Path(relative_path).parent.parts
    return folders


def preprocess_markdown(md_content):
    lines = md_content.split('\n')
    processed_lines = []
    
    for line in lines:
        if '```' in line and any(keyword in line for keyword in ['Polars', 'DuckDB', 'Pandas', 'PySpark', '极地', '鸭子数据库', '熊猫']):
            if line.strip().endswith('```'):
                processed_lines.append('```')
            else:
                continue
        else:
            processed_lines.append(line)
    
    md_content = '\n'.join(processed_lines)
    
    lines = md_content.split('\n')
    processed_lines = []
    
    for line in lines:
        stripped = line.lstrip()
        if stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.', stripped):
            indent = len(line) - len(line.lstrip())
            
            if indent == 2:
                line = '    ' + stripped
            elif indent == 3:
                line = '    ' + stripped
            elif indent == 4:
                line = '        ' + stripped
            elif indent == 5:
                line = '        ' + stripped
            elif indent == 6:
                line = '            ' + stripped
        
        processed_lines.append(line)
    
    return '\n'.join(processed_lines)


def parse_markdown_to_html(md_content):
    md_content = preprocess_markdown(md_content)
    return markdown.markdown(md_content, extensions=['tables', 'fenced_code'])


def reset_list_numbering(doc):
    pass


def add_list_item_to_doc(doc, li, level=0, is_ordered=False, counter=None):
    nested_ul = li.find('ul')
    nested_ol = li.find('ol')
    
    if nested_ul:
        nested_ul.extract()
    if nested_ol:
        nested_ol.extract()
    
    text = li.get_text().strip()
    
    if level == 0:
        if is_ordered:
            if counter is not None:
                counter[0] += 1
                p = doc.add_paragraph(f'{counter[0]}. {text}')
            else:
                p = doc.add_paragraph(f'1. {text}')
        else:
            p = doc.add_paragraph(text, style='List Bullet')
    else:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25 * level)
        if is_ordered:
            if counter is not None:
                counter[0] += 1
                run = p.add_run(f'{counter[0]}. {text}')
            else:
                run = p.add_run(f'1. {text}')
        else:
            run = p.add_run(f'• {text}')
    
    if nested_ul:
        for nested_li in nested_ul.find_all('li', recursive=False):
            add_list_item_to_doc(doc, nested_li, level=level + 1, is_ordered=False)
    
    if nested_ol:
        nested_counter = [0]
        for nested_li in nested_ol.find_all('li', recursive=False):
            add_list_item_to_doc(doc, nested_li, level=level + 1, is_ordered=True, counter=nested_counter)


def add_html_to_doc(doc, html_content, base_heading_level=1, md_file_path=None, docs_path=None):
    soup = BeautifulSoup(html_content, 'html.parser')
    
    processed_elements = set()
    
    for element in soup.find_all(True):
        if element in processed_elements:
            continue
        
        if element.name == 'h1':
            level = base_heading_level
            if level <= 9:
                p = doc.add_heading(element.get_text(), level=level)
            else:
                p = doc.add_paragraph(element.get_text())
                p.runs[0].bold = True
            reset_list_numbering(doc)
        elif element.name == 'h2':
            level = base_heading_level + 1
            if level <= 9:
                p = doc.add_heading(element.get_text(), level=level)
            else:
                p = doc.add_paragraph(element.get_text())
                p.runs[0].bold = True
            reset_list_numbering(doc)
        elif element.name == 'h3':
            level = base_heading_level + 2
            if level <= 9:
                p = doc.add_heading(element.get_text(), level=level)
            else:
                p = doc.add_paragraph(element.get_text())
                p.runs[0].bold = True
            reset_list_numbering(doc)
        elif element.name == 'h4':
            level = base_heading_level + 3
            if level <= 9:
                p = doc.add_heading(element.get_text(), level=level)
            else:
                p = doc.add_paragraph(element.get_text())
                p.runs[0].bold = True
            reset_list_numbering(doc)
        elif element.name == 'h5':
            level = base_heading_level + 4
            if level <= 9:
                p = doc.add_heading(element.get_text(), level=level)
            else:
                p = doc.add_paragraph(element.get_text())
                p.runs[0].bold = True
            reset_list_numbering(doc)
        elif element.name == 'h6':
            level = base_heading_level + 5
            if level <= 9:
                p = doc.add_heading(element.get_text(), level=level)
            else:
                p = doc.add_paragraph(element.get_text())
                p.runs[0].bold = True
            reset_list_numbering(doc)
        elif element.name == 'p':
            img = element.find('img')
            if img:
                img_src = img.get('src', '')
                if img_src:
                    img_path = resolve_image_path(img_src, md_file_path, docs_path)
                    if img_path and os.path.exists(img_path):
                        try:
                            p = doc.add_paragraph()
                            run = p.add_run()
                            run.add_picture(img_path, width=Inches(6))
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            p = doc.add_paragraph(f'[Image: {img_src}]')
                    else:
                        p = doc.add_paragraph(f'[Image not found: {img_src}]')
            else:
                text = element.get_text().strip()
                if text:
                    p = doc.add_paragraph(text)
        elif element.name == 'ul':
            processed_elements.update(element.find_all(True))
            for li in element.find_all('li', recursive=False):
                add_list_item_to_doc(doc, li, level=0, is_ordered=False)
        elif element.name == 'ol':
            processed_elements.update(element.find_all(True))
            counter = [0]
            for li in element.find_all('li', recursive=False):
                add_list_item_to_doc(doc, li, level=0, is_ordered=True, counter=counter)
        elif element.name == 'pre':
            code_element = element.find('code')
            if code_element:
                code_text = code_element.get_text()
            else:
                code_text = element.get_text()
            
            lines = code_text.split('\n')
            for line in lines:
                if line.strip():
                    cleaned_line = re.sub(r'^\d+', '', line)
                    p = doc.add_paragraph(cleaned_line, style='代码')
        elif element.name == 'blockquote':
            processed_elements.update(element.find_all(True))
            p = doc.add_paragraph(element.get_text())
            try:
                p.style = 'Intense Quote'
            except:
                pass
        elif element.name == 'table':
            processed_elements.update(element.find_all(True))
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
                table.style = 'Light Grid Accent 1'
                
                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for j, cell in enumerate(cells):
                        table.rows[i].cells[j].text = cell.get_text().strip()
                        if cell.name == 'th':
                            for run in table.rows[i].cells[j].paragraphs[0].runs:
                                run.bold = True


def resolve_image_path(img_src, md_file_path, docs_path):
    md_dir = os.path.dirname(md_file_path)
    
    if os.path.isabs(img_src):
        return img_src
    
    relative_path = os.path.join(md_dir, img_src)
    if os.path.exists(relative_path):
        return relative_path
    
    if docs_path:
        docs_relative_path = os.path.join(docs_path, img_src)
        if os.path.exists(docs_relative_path):
            return docs_relative_path
    
    return None


def convert_markdown_to_word(docs_path, output_path, template_path=None):
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
        for para in doc.paragraphs[:]:
            para._element.getparent().remove(para._element)
        for table in doc.tables[:]:
            table._element.getparent().remove(table._element)
    else:
        doc = Document()
        doc.add_heading('Documentation', level=0)
    
    md_files = []
    
    for root, dirs, files in os.walk(docs_path):
        for file in files:
            if file.endswith('.md'):
                file_path = os.path.join(root, file)
                md_files.append(file_path)
    
    md_files.sort(key=lambda x: os.path.getctime(x))
    
    created_folders = set()
    
    for md_file in md_files:
        folder_structure = get_folder_structure(md_file, docs_path)
        depth = len(folder_structure)
        base_heading_level = depth + 1
        
        for i, folder in enumerate(folder_structure):
            folder_path = '/'.join(folder_structure[:i+1])
            if folder_path not in created_folders:
                heading_level = i + 1
                if heading_level <= 9:
                    doc.add_heading(folder, level=heading_level)
                else:
                    p = doc.add_paragraph(folder)
                    p.runs[0].bold = True
                created_folders.add(folder_path)
        
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        file_name = Path(md_file).stem
        file_name = file_name.replace('-', ' ').replace('_', ' ')
        file_name = ' '.join(word.capitalize() for word in file_name.split())
        
        if base_heading_level <= 9:
            doc.add_heading(file_name, level=base_heading_level)
        else:
            p = doc.add_paragraph(file_name)
            p.runs[0].bold = True
        
        html_content = parse_markdown_to_html(md_content)
        add_html_to_doc(doc, html_content, base_heading_level=base_heading_level + 1, 
                       md_file_path=md_file, docs_path=docs_path)
        
        doc.add_paragraph()
    
    doc.save(output_path)
    print(f'Word document created successfully: {output_path}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert Markdown files to Word document')
    parser.add_argument('docs_path', help='Path to the folder containing Markdown files')
    parser.add_argument('output_path', help='Path for the output Word document')
    parser.add_argument('--template', help='Path to template Word document (optional)', default=None)
    
    args = parser.parse_args()
    
    docs_path = args.docs_path
    output_path = args.output_path
    template_path = args.template
    
    if not os.path.exists(docs_path):
        print(f'Error: The specified folder does not exist: {docs_path}')
        sys.exit(1)
    
    try:
        convert_markdown_to_word(docs_path, output_path, template_path)
    except Exception as e:
        print(f'Error: {str(e)}')
        sys.exit(1)
